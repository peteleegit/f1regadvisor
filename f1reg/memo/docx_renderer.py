"""Assembles MemoContext into a downloadable Word (.docx) document.

Structure mirrors renderer.py — verdict-first, appendix at the bottom.
Users can edit and annotate the output before sending to counsel.
"""
from __future__ import annotations

import datetime
import io
import re

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from f1reg.context import MemoContext


# ---------------------------------------------------------------------------
# Low-level helpers
# ---------------------------------------------------------------------------

def _runs(para, text: str) -> None:
    """Add runs to a paragraph, parsing **bold** and *italic* inline markup."""
    for part in re.split(r"(\*\*.*?\*\*|\*[^*]+\*)", text):
        if part.startswith("**") and part.endswith("**"):
            para.add_run(part[2:-2]).bold = True
        elif part.startswith("*") and part.endswith("*"):
            para.add_run(part[1:-1]).italic = True
        else:
            para.add_run(part)


def _hr(doc: Document) -> None:
    """Add a horizontal rule via paragraph bottom border."""
    para = doc.add_paragraph()
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "auto")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _bullets(doc: Document, markdown: str) -> None:
    """Add bullet list items from '- item' markdown lines."""
    for line in markdown.splitlines():
        line = line.strip()
        if line.startswith("- "):
            para = doc.add_paragraph(style="List Bullet")
            _runs(para, line[2:])


def _freeform(doc: Document, text: str) -> None:
    """Render free-form markdown text: headings, bullets, and body paragraphs."""
    for line in text.splitlines():
        s = line.strip()
        if not s:
            continue
        if s.startswith("### "):
            doc.add_heading(s[4:], 3)
        elif s.startswith("## "):
            doc.add_heading(s[3:], 2)
        elif s.startswith("# "):
            doc.add_heading(s[2:], 1)
        elif s.startswith("- ") or s.startswith("* "):
            para = doc.add_paragraph(style="List Bullet")
            _runs(para, s[2:])
        elif s.startswith("---"):
            _hr(doc)
        else:
            _runs(doc.add_paragraph(), s)


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def build_memo_docx(ctx: MemoContext) -> bytes:
    """Return .docx bytes for the memo described by ctx."""
    doc = Document()
    concept = ctx.concept
    date_str = datetime.date.today().strftime("%d %B %Y").lstrip("0")

    # ------------------------------------------------------------------
    # Header
    # ------------------------------------------------------------------
    doc.add_heading("F1 Regulatory Risk Memo", 0)

    for label, value in [
        ("Concept:", concept.summary),
        ("Season:", concept.season),
        ("Assessed:", date_str),
    ]:
        p = doc.add_paragraph()
        p.add_run(label + " ").bold = True
        p.add_run(value)

    _hr(doc)

    # ------------------------------------------------------------------
    # Bottom Line
    # ------------------------------------------------------------------
    doc.add_heading("Bottom Line", 1)
    verdict = ctx.bottom_line_verdict or "ASSESSMENT INCOMPLETE"
    confidence = ctx.bottom_line_confidence or ""
    p = doc.add_paragraph()
    p.add_run(verdict + (f" — {confidence}" if confidence else "")).bold = True
    if ctx.bottom_line_summary:
        _runs(doc.add_paragraph(), ctx.bottom_line_summary)

    # ------------------------------------------------------------------
    # Recommended Action
    # ------------------------------------------------------------------
    doc.add_heading("Recommended Action", 1)
    action = ctx.recommended_action or "(not assessed)"
    rationale = ctx.recommended_action_rationale or ""
    p = doc.add_paragraph()
    p.add_run(action + ".").bold = True
    if rationale:
        p.add_run("  " + rationale)

    _hr(doc)

    # ------------------------------------------------------------------
    # Arguments
    # ------------------------------------------------------------------
    if ctx.arguments_for:
        doc.add_heading("Arguments For Legality", 1)
        doc.add_paragraph().add_run("[Advocacy — not factual synthesis]").italic = True
        _bullets(doc, ctx.arguments_for)

    if ctx.arguments_against:
        doc.add_heading("Arguments Against / Key Risks", 1)
        doc.add_paragraph().add_run("[Advocacy — not factual synthesis]").italic = True
        _bullets(doc, ctx.arguments_against)

    # ------------------------------------------------------------------
    # Political and Protest Risk
    # ------------------------------------------------------------------
    if ctx.political_risk_summary:
        doc.add_heading("Political and Protest Risk", 1)
        _runs(doc.add_paragraph(), ctx.political_risk_summary)

    # ------------------------------------------------------------------
    # Mitigations
    # ------------------------------------------------------------------
    if ctx.mitigations:
        doc.add_heading("Mitigations", 1)
        _bullets(doc, ctx.mitigations)

    # ------------------------------------------------------------------
    # Open Questions
    # ------------------------------------------------------------------
    oq_parts: list[str] = []
    if ctx.open_questions:
        oq_parts.append(ctx.open_questions)
    low_findings = [f for f in ctx.audit_findings if f.severity == "LOW"]
    for f in low_findings:
        oq_parts.append(f"- [Audit flag — {f.implicated_agent}] {f.issue} {f.recommendation}")

    if oq_parts:
        doc.add_heading("Open Questions", 1)
        for part in oq_parts:
            _bullets(doc, part)

    # ------------------------------------------------------------------
    # HIGH audit flags
    # ------------------------------------------------------------------
    high_findings = [f for f in ctx.audit_findings if f.severity == "HIGH"]
    if high_findings:
        doc.add_heading("⚠ AUDIT FLAGS — review before acting", 1)
        for f in high_findings:
            p = doc.add_paragraph()
            p.add_run(f"{f.implicated_agent}: ").bold = True
            p.add_run(f.issue)
            p2 = doc.add_paragraph()
            p2.add_run(f"Recommendation: {f.recommendation}").italic = True

    # ------------------------------------------------------------------
    # Disclaimer
    # ------------------------------------------------------------------
    _hr(doc)
    doc.add_paragraph().add_run(
        "This memo is AI-generated by F1RegAdvisor. All conclusions must be verified "
        "by qualified regulatory counsel before any regulatory decision is made."
    ).italic = True

    # ------------------------------------------------------------------
    # Appendix — full analysis
    # ------------------------------------------------------------------
    _hr(doc)
    doc.add_heading("Appendix — Full Analysis", 1)

    if ctx.regulatory_summary:
        doc.add_heading("Regulatory Analysis", 2)
        _freeform(doc, ctx.regulatory_summary)
        rto = ctx.regulatory_text_output
        if rto and rto.controlling_articles:
            doc.add_paragraph().add_run("Controlling articles:").bold = True
            for a in rto.controlling_articles:
                doc.add_paragraph(a, style="List Bullet")
        if rto and rto.gaps:
            doc.add_paragraph().add_run("Regulatory gaps:").bold = True
            for g in rto.gaps:
                doc.add_paragraph(g, style="List Bullet")

    if ctx.precedent_summary:
        doc.add_heading("Precedent Analysis", 2)
        _freeform(doc, ctx.precedent_summary)
        po = ctx.precedent_output
        if po and po.comparable_cases:
            doc.add_paragraph().add_run("Comparable cases:").bold = True
            for c in po.comparable_cases:
                doc.add_paragraph(c, style="List Bullet")

    if ctx.political_economy_analysis:
        doc.add_heading("Political Economy (Full)", 2)
        _freeform(doc, ctx.political_economy_analysis)

    if ctx.rival_protest_analysis:
        doc.add_heading("Rival Protest Simulation", 2)
        _freeform(doc, ctx.rival_protest_analysis)

    if ctx.debate_rounds:
        _round_labels = {
            1: "Round 1: Opening Statements",
            2: "Round 2: Rebuttals",
            3: "Round 3: Closing Statements",
        }
        doc.add_heading("Debate Transcript", 2)
        for r in ctx.debate_rounds:
            doc.add_heading(_round_labels.get(r.round_number, f"Round {r.round_number}"), 3)
            doc.add_paragraph().add_run("FIA Skeptic:").bold = True
            _freeform(doc, r.skeptic_argument)
            doc.add_paragraph().add_run("Liberal Construction / Defence:").bold = True
            _freeform(doc, r.defense_argument)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
